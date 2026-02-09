import os
import sys
import logging
from typing import List
import uuid
from fastapi import APIRouter, HTTPException, Request, Response, File, UploadFile
from fastapi.responses import StreamingResponse
import json

# from com.sequation.document.service.azureTableService import AzureTableService
from dotenv import load_dotenv
import datetime
from azure.data.tables import TableServiceClient
from service.rag import run_rag
from service.lawyer_rag import run_lawyer_rag
from service.file_utils import extract_text_from_bytes
from service.format_utils import format_timestamp, get_next_case_number
import docx
from io import BytesIO

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()
load_dotenv()
connection_string = os.environ.get("AZURE_STORAGE_CONNECTION_STRING")
table_service = TableServiceClient.from_connection_string(connection_string)
aila_log_table = table_service.get_table_client("ailalogs")
case_status_table = table_service.get_table_client("ailacasestatus")


@router.post('/export_analysis')
def export_analysis(req: Request) -> Response:
    try:
        # Get the analysis data from request body
        analysis = req.json()
        
        # Create a new Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading('AILA Case Analysis Report', 0)
        
        # Add date and time
        current_time = datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")
        doc.add_paragraph(f"Generated on {current_time}")
        
        # Add parties section
        doc.add_heading('Parties', level=1)
        for party in analysis.get('parties', []):
            doc.add_paragraph(
                f"{party['name']} - {party['role']}",
                style='List Bullet'
            )
        
        # Add facts section
        doc.add_heading('Facts', level=1)
        for fact in analysis.get('facts', []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{fact['status']}] ").bold = True
            p.add_run(fact['fact'])
            
        # Add suggested rulings section
        doc.add_heading('Analysis', level=1)
        for i, ruling in enumerate(analysis.get('suggested_rulings', []), 1):
            # Add issue heading
            doc.add_heading(f"Issue {i}: {ruling['issue']}", level=2)
            
            # Add evidence
            doc.add_heading('Evidence', level=3)
            doc.add_paragraph(ruling['evidence'])
            
            # Add relevant articles
            if ruling.get('relevant_articles'):
                doc.add_heading('Relevant Articles', level=3)
                for article in ruling['relevant_articles']:
                    # Article title
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"Article {article['article_number']} - {article['legislation_title']}").bold = True
                    
                    # Article quote
                    doc.add_paragraph(f"Quote: {article['article_quote']}")
                    
                    # Article explanation
                    doc.add_paragraph(f"Explanation: {article['explanation']}")
                    
                    if article.get('full_article_text'):
                        doc.add_paragraph(f"Full Text: {article['full_article_text']}")
            
            # Add ruling
            doc.add_heading('Ruling', level=3)
            doc.add_paragraph(ruling['suggested_ruling'])
            
        # Add final ruling if present
        if analysis.get('final_ruling'):
            doc.add_heading('Final Ruling', level=1)
            doc.add_paragraph(analysis['final_ruling'])
            
        # Add final court orders if present
        if analysis.get('final_court_orders'):
            doc.add_heading('Final Court Orders', level=1)
            for order in analysis['final_court_orders']:
                doc.add_paragraph(order, style='List Number')
        
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response with the document
        return Response(
            body=doc_io.getvalue(),
            status_code=200,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                'Content-Disposition': 'attachment; filename=case_analysis.docx'
            }
        )
        
    except Exception as e:
        logging.error(f"[API ERROR][export_analysis] Error generating document: {str(e)}")
        return Response(
            "Error generating document",
            status_code=500
        )