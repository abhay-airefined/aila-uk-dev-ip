import uuid
import re
import logging
from typing import Tuple, Optional, Any, List
from datetime import datetime
from dotenv import load_dotenv
import requests
from azure.storage.blob import BlobServiceClient
from urllib.parse import urljoin
from fastapi import HTTPException
import os
import weaviate
from weaviate.classes import config as Config
from weaviate.classes import query as Query
from weaviate.collections.classes.config_vector_index import _VectorIndexConfigCreate as VectorIndexConfig
from weaviate.collections.classes.config_base import _QuantizerConfigCreate as QuantizerConfig
from io import BytesIO
import docx
from pypdf import PdfReader
from service.config import Config as ServiceConfig
from weaviate.classes.query import MetadataQuery
from .policy_parser import PolicyBenefitParser
from service.splitter import SuperRecursiveSplitter
from datetime import datetime, time
import hashlib

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
client = ServiceConfig.buildWeaviateConnection()
weaviate_collection_name = os.getenv("WEAVIATE_COLLECTION_NAME")
load_dotenv()

class WeaviateService:
    def __init__(self):
        logger.info("Initializing WeaviateService")
        
        self.connection_string = os.getenv("CONNECTION_STRING")
        self.blob_service_client = BlobServiceClient.from_connection_string(self.connection_string)
        self.container_client = self.blob_service_client.get_container_client(os.getenv("CONTAINER_NAME"))
        self.api_key = os.getenv("AZURE_OPENAI_EMBEDDING_API_KEY")
        self.azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        self.api_version = os.getenv("AZURE_OPENAI_EMBEDDING_VERSION")
        self.resource_name = os.getenv("AZURE_OPENAI_RESOURCE_NAME")
        self.deployment_id = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
        self.policy_benefit_collection_name = os.getenv("WEAVIATE_COLLECTION_NAME")
        
        logger.info(f"WeaviateService config - collection: {self.policy_benefit_collection_name}, endpoint: {self.azure_endpoint}")
        
        logger.info("Building Weaviate connection")
        self.client = None
        self._ensure_connection()
        
        logger.info("Ensuring collection exists")
        self.create_policy_benefit_class(self.client)
        
        logger.info("WeaviateService initialization completed")

    def _ensure_connection(self):
        """Ensure Weaviate client is connected, reconnect if necessary"""
        try:
            if self.client is None:
                logger.info("Creating new Weaviate client connection")
                self.client = ServiceConfig.buildWeaviateConnection()
                return
            
            # Check if client is still connected
            if not self.client.is_live():
                logger.warning("Weaviate client connection lost. Reconnecting...")
                self.client = ServiceConfig.buildWeaviateConnection()
                logger.info("Weaviate client reconnected successfully")
        except Exception as e:
            logger.error(f"Error ensuring Weaviate connection: {str(e)}", exc_info=True)
            raise

    def create_policy_benefit_class(self, client: weaviate.WeaviateClient):
        """Creates the unified 'PolicyBenefit' collection in Weaviate if it doesn't exist."""
        logger.info(f"Checking for collection: {self.policy_benefit_collection_name}")
        self._ensure_connection()
        
        if not client.collections.exists(self.policy_benefit_collection_name):
            logger.info(f"Collection {self.policy_benefit_collection_name} does not exist, creating it")
            client.collections.create(
                name=self.policy_benefit_collection_name,
                description="Unified collection for all document types",
                vectorizer_config=[
                    Config.Configure.NamedVectors.text2vec_azure_openai(
                        name="default",
                        source_properties=["rawText", "caseName"], # vectorize the raw text
                        base_url=self.azure_endpoint,
                        resource_name=self.resource_name,
                        deployment_id=self.deployment_id,
                        vector_index_config=Config.Configure.VectorIndex.hnsw(
                            quantizer=Config.Configure.VectorIndex.Quantizer.bq(rescore_limit=200)
                        )
                    )
                ],
                generative_config=Config.Configure.Generative.azure_openai(
                    resource_name=self.resource_name,
                    deployment_id=self.deployment_id
                ),
                multi_tenancy_config=Config.Configure.multi_tenancy(enabled=True),
                properties=[
                    Config.Property(name="type", data_type=Config.DataType.TEXT),
                    Config.Property(name="parentDocumentId", data_type=Config.DataType.TEXT),
                    Config.Property(name="rawText", data_type=Config.DataType.TEXT),
                    Config.Property(name="caseName", data_type=Config.DataType.TEXT),
                    Config.Property(name="caseDetails", data_type=Config.DataType.TEXT),
                    Config.Property(name="from", data_type=Config.DataType.TEXT_ARRAY),
                    Config.Property(name="publishedDate", data_type=Config.DataType.DATE),
                    Config.Property(name="category", data_type=Config.DataType.TEXT_ARRAY),
                    Config.Property(name="subCategory", data_type=Config.DataType.TEXT_ARRAY),
                    Config.Property(name="landmark", data_type=Config.DataType.TEXT),
                    Config.Property(name="decisionDate", data_type=Config.DataType.DATE),
                    Config.Property(name="country", data_type=Config.DataType.TEXT),
                    Config.Property(name="jurisdictionCode", data_type=Config.DataType.TEXT),
                    Config.Property(name="blobUrl", data_type=Config.DataType.TEXT),
                    Config.Property(name="md5Hash", data_type=Config.DataType.TEXT),
                    Config.Property(name="orgId", data_type=Config.DataType.INT)
                ]
            )
            logger.info(f"Collection {self.policy_benefit_collection_name} created successfully")
        else:
            logger.info(f"Collection {self.policy_benefit_collection_name} already exists")

    @staticmethod
    def process_document(fileUrl: str, contentType: str) -> Tuple[str, Optional[List[str]]]:
        """
        Process a document and return its text content and pages if applicable.
        """
        logger.info(f"Processing document - URL: {fileUrl[:50]}..., content type: {contentType}")
        
        if not fileUrl:
            logger.error("No file content received")
            raise ValueError("No file content received")

        text = ""
        pages = None # Pages are only relevant for generic chunking now
        
        try:
            if 'text/markdown' in contentType:
                logger.info("Processing markdown document")
                text = fileUrl
                pages = [text]
                
                logger.info("Markdown processed successfully")
            elif 'application/pdf' in contentType:    
                logger.info("Processing PDF document")
                response = requests.get(fileUrl)
                response.raise_for_status()
                pdf_file = BytesIO(response.content)
                reader = PdfReader(pdf_file)
                pages = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    pages.append(page_text)
                    text += page_text
                logger.info(f"PDF processed successfully. Pages: {len(pages)}")
            elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in contentType:
                logger.info("Processing DOCX document")
                response = requests.get(fileUrl)
                response.raise_for_status()
                docx_file = BytesIO(response.content)
                doc = docx.Document(docx_file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                logger.info(f"DOCX processed successfully. Paragraphs: {len(doc.paragraphs)}")
            elif 'text/plain' in contentType:
                logger.info("Processing plain text document")
                response = requests.get(fileUrl)
                response.raise_for_status()
                text = response.text
                logger.info("Plain text processed successfully")
            else:
                logger.error(f"Unsupported file type: {contentType}")
                raise ValueError(f"Unsupported file type: {contentType}")

            if not text.strip():
                logger.error("Could not extract text from the document")
                raise ValueError("Could not extract text from the document")
            
            logger.info(f"Document processing completed. Text length: {len(text)}")
            return text.strip(), pages

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            raise

    def create_policy_benefit_chunks(self, markdown_text: str, documentUrl: str, filename: str, orgId: int, md5Hash: str) -> List[dict]:
        """
        Processes a markdown document to extract structured policy benefit chunks.
        """
        try:
            parser = PolicyBenefitParser(
                markdown_text=markdown_text,
                document_url=documentUrl,
                filename=filename,
                org_id=orgId,
                md5_hash=md5Hash
            )
            return parser.parse()
        except Exception as e:
            logger.error(f"Failed to create policy benefit chunks for {filename}: {e}", exc_info=True)
            return []

    def upload_documents(self, chunks: List[dict], tenant_name: str) -> None:
        """
        Uploads a batch of structured document objects to the unified 'PolicyBenefit' collection.
        """
        logger.info(f"Uploading {len(chunks)} document chunks to '{self.policy_benefit_collection_name}' for tenant: {tenant_name}")
        try:
            self._ensure_connection()
            collection = self.client.collections.get(self.policy_benefit_collection_name)

            if not collection.tenants.exists(tenant_name):
                logger.info(f"Creating new tenant: {tenant_name} in collection {self.policy_benefit_collection_name}")
                collection.tenants.create([weaviate.classes.tenants.Tenant(name=tenant_name)])
            
            tenant_collection = collection.with_tenant(tenant_name)

            with tenant_collection.batch.dynamic() as batch:
                for chunk_properties in chunks:
                    batch.add_object(
                        properties=chunk_properties,
                        uuid=uuid.uuid4()
                    )
            failed_objects = tenant_collection.batch.failed_objects
            if failed_objects:
                logger.error(f"Number of failed imports: {len(failed_objects)}")
                logger.error(f"First failed object: {failed_objects[0]}")

                md5_hashes = [chunk.get("md5Hash") for chunk in chunks if chunk.get("md5Hash")]

                if md5_hashes:
                    logger.info(f"Deleting failed batch using md5Hashes: {md5_hashes}")
                    tenant_collection.data.delete_many(
                        where=weaviate.classes.query.Filter.by_property("md5Hash").contains_any(md5_hashes)
                    )
                    logger.info("Successfully deleted failed batch.")
            
            logger.info(f"Successfully uploaded batch for {len(chunks)} document chunks.")

        except Exception as e:
            logger.error(f"Error uploading document chunks: {str(e)}", exc_info=True)
            raise

    def search_documents(self, query: str, top_k: int, tenant_name: str) -> List[dict]:
        """
        Search for relevant policy benefits using semantic similarity.
        This method is now configured to search the structured 'PolicyBenefit' collection.
        """
        logger.info(f"Searching for policy benefits - query: '{query[:50]}...', top_k: {top_k}, tenant: {tenant_name}")
        try:
            self._ensure_connection()
            
            policy_collection = self.client.collections.get(self.policy_benefit_collection_name)
            tenant_collection = policy_collection.with_tenant(tenant_name)
            
            logger.info(f"Performing hybrid search in '{self.policy_benefit_collection_name}' collection...")
                        
            response = tenant_collection.query.hybrid(
                query=query,
                limit=top_k,
                alpha=0.7,
                query_properties=["rawText", "notes", "section"],
                return_metadata=MetadataQuery(score=True)
            )
            
            logger.info(f"Search completed. Found {len(response.objects)} potential benefit results.")
            
            # Format results with the rich, structured data
            results = []
            for obj in response.objects:
                properties = obj.properties
                result = {
                    "rawText": properties.get("rawText"),
                    "section": properties.get("section"),
                    "title": properties.get("title"),
                    "description": properties.get("description"),
                    "notes": properties.get("notes"),
                    "filename": properties.get("filename"),
                    "coverage_network": properties.get("coverage_network"),
                    "coverage_nonNetwork": properties.get("coverage_nonNetwork")
                }
                results.append(result)

            logger.info(f"Policy benefit search completed successfully. Returning {len(results)} structured results.")
            return results
            
        except Exception as e:
            logger.error(f"Error searching for policy benefits: {str(e)}", exc_info=True)
            raise
        
    def check_document_exists(self, md5Hash: str, tenant_name: str, orgId: int) -> bool:
        """
        Check if a document exists in the unified collection for the given tenant.
        """
        logger.info(f"Checking document existence - md5Hash: {md5Hash[:20]}..., tenant: {tenant_name}, org: {orgId}")
        try:
            self._ensure_connection()
            collection = self.client.collections.get(self.policy_benefit_collection_name)
            
            if not collection.tenants.exists(tenant_name):
                logger.info(f"Tenant {tenant_name} does not exist in {self.policy_benefit_collection_name}. Concluding document does not exist.")
                return False

            tenant_collection = collection.with_tenant(tenant_name)
            
            response = tenant_collection.query.fetch_objects(
                limit=1,
                filters=weaviate.classes.query.Filter.by_property("md5Hash").equal(md5Hash)
            )
            
            exists = len(response.objects) > 0
            if exists:
                logger.info(f"Document with hash {md5Hash} found in collection '{self.policy_benefit_collection_name}' for tenant '{tenant_name}'.")
            return exists

        except Exception as e:
            logger.warning(f"An error occurred while checking existence in '{self.policy_benefit_collection_name}'. Assuming document does not exist. Error: {e}")
            return False
        
    def chunk_and_embed_document(self, full_text: str, doc_type: str, doc, full_pdf_blob_path: str, md5Hash: str, orgId: int) -> List[dict]:
        """
        Chunk and embed a document, handling pages if present, using parallel processing.
        """
        logger.info(f"Starting chunk and embed document")

        doc_splitter = SuperRecursiveSplitter(
            separators=["\n\n", "\n", ".", ",", " "],
            target_chunk_size=20000,
            separator_placeholders=True,
            overlap=0,
            reconstruct=True,
            verbosity=0
        )
        
        logger.info("Splitting document into chunks")
        chunks = doc_splitter.split_into_chunks(full_text)
        logger.info(f"Document split into {len(chunks)} chunks")
                
        # Process chunks in batches
        BATCH_SIZE = 8
        embedded_chunks = []
        
        logger.info(f"Processing {len(chunks)} chunks in batches of {BATCH_SIZE}")
        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i:i + BATCH_SIZE]
            batch_num = (i // BATCH_SIZE) + 1
            logger.info(f"Processing batch {batch_num} with {len(batch)} chunks")
            
            try:                
                for chunk in batch:                    
                    if doc_type == 'decision':
                        chunk_obj = {
                            "type": "decision",
                            "parentDocumentId": doc.id,
                            "rawText": chunk,
                            "caseName": doc.case_name,
                            "from": doc.from_tribunal if doc.from_tribunal else [],
                            "publishedDate": datetime.combine(doc.published_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.published_date else None,
                            "country": doc.country if doc.country else "",
                            "jurisdictionCode": doc.jurisdiction_code if doc.jurisdiction_code else "",
                            "decisionDate": datetime.combine(doc.decision_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.decision_date else None,
                            "blobUrl": f"{full_pdf_blob_path}",
                            "md5Hash": md5Hash,
                            "orgId": orgId,
                        }
                    else:
                        chunk_obj = {
                            "type": "appeal",
                            "parentDocumentId": doc.id,
                            "rawText": chunk,
                            "caseName": doc.case_name,
                            "caseDetails": doc.case_details,
                            "from": doc.from_tribunal if doc.from_tribunal else [],
                            "publishedDate": datetime.combine(doc.published_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.published_date else None,
                            "category": doc.category if doc.category else [],
                            "subCategory": doc.sub_category if doc.sub_category else [],
                            "landmark": doc.landmark if doc.landmark else "",
                            "decisionDate": datetime.combine(doc.decision_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.decision_date else None,
                            "blobUrl": f"{full_pdf_blob_path}",
                            "md5Hash": md5Hash,
                            "orgId": orgId,
                            }
                   
                    embedded_chunks.append(chunk_obj)
                                    
            except Exception as e:
                logger.error(f"Error processing batch {batch_num}: {str(e)}", exc_info=True)
                raise
        
        logger.info(f"Document chunking completed")
        return embedded_chunks

    def create_document_chunk(self, full_text: str, doc_type: str, doc, full_pdf_blob_path: str, md5Hash: str, orgId: int) -> List[dict]:
        logger.info(f"Creating a single document chunk for file: {doc.id}")
        
        # md5Hash = hashlib.md5(full_text.encode('UTF-8')).hexdigest()
        
        if doc_type == 'decision':
            chunk_obj = {
                "type": "decision",
                "parentDocumentId": doc.id,
                "rawText": full_text,
                "caseName": doc.case_name,
                "from": doc.from_tribunal if doc.from_tribunal else [],
                "publishedDate": datetime.combine(doc.published_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.published_date else None,
                "country": doc.country if doc.country else "",
                "jurisdictionCode": doc.jurisdiction_code if doc.jurisdiction_code else "",
                "decisionDate": datetime.combine(doc.decision_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.decision_date else None,
                "blobUrl": f"{full_pdf_blob_path}",
                "md5Hash": md5Hash,
                "orgId": orgId,
            }
        else:
            chunk_obj = {
                "type": "appeal",
                "parentDocumentId": doc.id,
                "rawText": full_text,
                "caseName": doc.case_name,
                "caseDetails": doc.case_details,
                "from": doc.from_tribunal if doc.from_tribunal else [],
                "publishedDate": datetime.combine(doc.published_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.published_date else None,
                "category": doc.category if doc.category else [],
                "subCategory": doc.sub_category if doc.sub_category else [],
                "landmark": doc.landmark if doc.landmark else "",
                "decisionDate": datetime.combine(doc.decision_date, time.min).strftime('%Y-%m-%dT%H:%M:%SZ') if doc.decision_date else None,
                "blobUrl": f"{full_pdf_blob_path}",
                "md5Hash": md5Hash,
                "orgId": orgId,
            }
        
        return [chunk_obj]


    def close(self):
        logger.info("Closing Weaviate client connection")
        if self.client:
            try:
                self.client.close()
                logger.info("Weaviate client connection closed successfully")
            except Exception as e:
                logger.warning(f"Error closing Weaviate client: {str(e)}")
        self.client = None

    def delete_collection(self):
        logger.info(f"Deleting collection: {self.policy_benefit_collection_name}")
        self._ensure_connection()
        self.client.collections.delete(self.policy_benefit_collection_name)
        logger.info("Collection deleted successfully")

    
    def search_relevant_docs(self, query: str, page: int, page_size: int) -> dict:
        """
        Find top 300 most relevant chunks for a query and paginate through them.
        """

        policy_collection = client.collections.get(weaviate_collection_name)
        tenant_collection = policy_collection.with_tenant("1")

        try:
            # Fetch top 300 most relevant records
            TOP_RESULTS_LIMIT = 300
            
            logger.info(f"Fetching top {TOP_RESULTS_LIMIT} most relevant records...")
            
            response = tenant_collection.query.hybrid(
                query=query,
                alpha=0.7,
                query_properties=["rawText", "caseName", "jurisdictionCode"],
                limit=TOP_RESULTS_LIMIT,  # Get exactly top 300
            )
            
            logger.info(f"Retrieved {len(response.objects)} objects from Weaviate")
            
            # Process all fetched results
            all_results = []
            for obj in response.objects:
                properties = obj.properties
                                
                result = {
                    "rawText": properties.get("rawText"),
                    "caseName": properties.get("caseName"),
                    "jurisdictionCode": properties.get("jurisdictionCode"),
                    "pdfBlobUrl": properties.get("blobUrl"),
                }
                all_results.append(result)

            # # Sort by relevance score (highest first) - Weaviate should already be sorted
            # all_results.sort(key=lambda x: x["relevanceScore"], reverse=True)
            
            # # Remove relevance scores from results if you don't want to expose them
            # for result in all_results:
            #     result.pop("relevanceScore", None)

            # Apply pagination to the top 300 results
            total_records = len(all_results)  # This will be <= 300
            total_pages = (total_records + page_size - 1) // page_size if total_records > 0 else 0
            
            start_index = (page - 1) * page_size
            end_index = start_index + page_size

            # Check if requested page is out of range
            if start_index >= total_records and total_records > 0:
                raise HTTPException(
                    status_code=404, 
                    detail=f"Page {page} is out of range. Maximum page is {total_pages}"
                )

            # Get the paginated results
            paginated_results = all_results[start_index:end_index]

            return {
                "query": query[:100] + "..." if len(query) > 100 else query,
                "page": page,
                "page_size": page_size,
                "total_records": total_records,  # Actual number of matching records (≤ 300)
                # "total_pages": total_pages,
                # "max_results_limit": TOP_RESULTS_LIMIT,  # Inform user about the 300 limit
                # "has_next": page < total_pages,
                # "has_previous": page > 1,
                # "results_count": len(paginated_results),
                # "is_limited_to_top_results": total_records == TOP_RESULTS_LIMIT,  # Flag if we hit the 300 limit
                "results": paginated_results
            }

        except Exception as e:
            logger.info(f"Error in search_relevant_docs: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")
